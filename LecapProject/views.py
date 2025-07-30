import requests
import io
import logging
from django.contrib.auth.decorators import login_required
from django.contrib.auth import get_user_model
from django.contrib import messages
from django.shortcuts import render, redirect
from django.conf import settings
from django.utils import timezone
from django import forms
from accounts.models import AdminSettings, KaitenUserRoleOverride
from .models import  ProjectRate, DefaultRoleRate
from accounts.forms import AdminSettingsForm, CustomUserForm
from LecapProject.kaiten_api import fetch_kaiten_roles, fetch_kaiten_projects, fetch_kaiten_cards, \
        fetch_kaiten_time_logs, fetch_kaiten_boards, fetch_kaiten_board_roles
from django.forms import modelformset_factory, ModelForm, HiddenInput
from datetime import datetime, timedelta
import pytz
from django.http import HttpResponse, JsonResponse
from docx import Document
from docxTemplate.models import TemplateFile
from docxTemplate.views import insert_table_after, set_table_borders, insert_paragraph_after_table, convert_number_to_text
from django.forms import ModelForm, HiddenInput
from .models import ProjectRate
from .forms import DefaultRoleRateFormSet

from docx.shared import Inches, Pt
from django.db.models.functions import Cast
from django.db.models import CharField
from urllib.parse import urlencode
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .kaiten_api import fetch_kaiten_roles, fetch_kaiten_users
from django.urls import reverse

logger = logging.getLogger('kaiten')

User = get_user_model()

@login_required
def templates_view(request):
    return render(request, 'templates.html')

@login_required
def administration_view(request):
    # Если пользователь — суперпользователь, можно сразу редиректить в панель администратора
    """if request.user.is_superuser:
        return redirect('admin:index')"""
    # Иначе можно отобразить кастомную страницу редактирования пользователей
    return render(request, 'custom_administration.html')

def check_board_rates(board, project_id, roles):
    """
    Функция проверяет, удовлетворяет ли доска условию наличия ставок для каждой роли.
    Если для роли отсутствует кастомная ставка, ищется дефолтная ставка.
    Возвращает:
      - board_valid: True, если для всех ролей найдена ставка (кастомная или дефолтная);
      - auto_used: True, если хотя бы для одной роли ставка берётся из дефолтных.
    """
    board_valid = True
    auto_used = False
    board_id = board.get("id")
    for role in roles:
        role_id = str(role.get("id"))
        try:
            pr = ProjectRate.objects.get(
                project_id=project_id,
                board_id=board_id,
                role_id=role_id
            )
            if pr.rate is None:
                try:
                    dr = DefaultRoleRate.objects.get(role_id=role_id)
                    if dr.default_rate is not None:
                        auto_used = True
                    else:
                        board_valid = False
                        break
                except DefaultRoleRate.DoesNotExist:
                    board_valid = False
                    break
        except ProjectRate.DoesNotExist:
            try:
                dr = DefaultRoleRate.objects.get(role_id=role_id)
                if dr.default_rate is not None:
                    auto_used = True
                else:
                    board_valid = False
                    break
            except DefaultRoleRate.DoesNotExist:
                board_valid = False
                break
    return board_valid, auto_used

@login_required
def get_boards(request):
    kaiten_api_down = False
    space_id = request.GET.get('space_id')
    for_report = request.GET.get('for_report') == "1"
    admin_settings, _ = AdminSettings.objects.get_or_create(pk=1)
    domain = admin_settings.url_domain_value_id
    bearer_key = admin_settings.api_auth_key

    boards = fetch_kaiten_boards(domain, bearer_key, space_id)
    if not boards:
        kaiten_api_down = True

    if for_report:
        global_roles = fetch_kaiten_roles(domain, bearer_key)
        for role in global_roles:
            DefaultRoleRate.objects.get_or_create(
                role_id=str(role.get('id')),
                defaults={'role_name': role.get('name')}
            )

        for board in boards:
            board_roles = fetch_kaiten_board_roles(domain, bearer_key, space_id, board['id'])
            if not board_roles:
                kaiten_api_down = True
            valid, auto_used = check_board_rates(board, space_id, board_roles)
            board["has_rates"] = valid
            if valid and auto_used:
                board["title"] += " (автоставки)"
    
    # Для AJAX‑ответа можно вернуть ошибку прямо в JSON
    error_message = ""
    if kaiten_api_down:
        error_message = "Сервер Kaiten недоступен. Пожалуйста, повторите попытку позже."
    return JsonResponse({"boards": boards, "error": error_message})

@login_required
def rates_view(request):
    admin_settings, _ = AdminSettings.objects.get_or_create(pk=1)
    domain = admin_settings.url_domain_value_id
    bearer_key = admin_settings.api_auth_key
    kaiten_api_down = False

    projects = []
    if domain and bearer_key:
        projects = fetch_kaiten_projects(domain, bearer_key)
        if not projects:
            kaiten_api_down = True

    project_id = request.GET.get('project_id')
    project_title = request.GET.get('project_title', '')
    board_id = request.GET.get('board_id')
    board_title = request.GET.get('board_title', '')

    if project_id and not project_title:
        for p in projects:
            if str(p['id']) == project_id:
                project_title = p['title']
                break
    if not project_id and projects:
        project_id = str(projects[0]['id'])
        project_title = projects[0]['title']

    boards = []
    if project_id:
        boards = fetch_kaiten_boards(domain, bearer_key, project_id)
        if not boards:
            kaiten_api_down = True
        if boards and not board_id:
            board_id = boards[0]['id']
            board_title = boards[0]['title']

    if domain and bearer_key and project_id and board_id:
        global_roles = fetch_kaiten_roles(domain, bearer_key)
        if not global_roles:
            kaiten_api_down = True
        else:
            for role in global_roles:
                role_id = str(role.get('id'))
                role_name = role.get('name')
                ProjectRate.objects.get_or_create(
                    project_id=project_id,
                    board_id=board_id,
                    role_id=role_id,
                    defaults={'role_name': role_name}
                )
                DefaultRoleRate.objects.get_or_create(
                    role_id=role_id,
                    defaults={'role_name': role_name}
                )
            #DefaultRoleRate.objects.exclude(role_id__in=api_role_ids).delete()

    ProjectRateFormSet = modelformset_factory(ProjectRate, form=ProjectRateForm, extra=0)
    rates_formset = None
    if project_id and board_id:
        rates_formset = ProjectRateFormSet(
            queryset=ProjectRate.objects.filter(project_id=project_id, board_id=board_id)
        )

    if kaiten_api_down:
        messages.error(request, "Сервер Kaiten недоступен. Пожалуйста, повторите попытку позже.")

    if request.method == 'POST':
        if 'save_default_rates' in request.POST:
            default_rate_formset = DefaultRoleRateFormSet(
                request.POST,
                queryset=DefaultRoleRate.objects.all()
            )
            if default_rate_formset.is_valid():
                default_rate_formset.save()
                messages.success(request, "Стандартные ставки сохранены.")
            else:
                messages.error(request, "Проверьте введённые данные в стандартных ставках.")
            params = {
                'project_id': project_id,
                'project_title': project_title,
                'board_id': board_id,
                'board_title': board_title,
            }
            return redirect(f"{reverse('rates')}?{urlencode(params)}")
        else:
            ProjectRateFormSet = modelformset_factory(ProjectRate, form=ProjectRateForm, extra=0)
            bound_formset = ProjectRateFormSet(
                request.POST,
                queryset=ProjectRate.objects.filter(
                    project_id=project_id,
                    board_id=board_id
                )
            )
            return save_rates(
                request,
                bound_formset,
                project_id,
                project_title,
                board_id,
                board_title
            )

    context = {
        'projects': projects,
        'boards': boards,
        'rates_formset': rates_formset,
        'default_rate_formset': DefaultRoleRateFormSet(
            queryset=DefaultRoleRate.objects.all()
        ),
        'selected_project_id': project_id,
        'selected_project_title': project_title,
        'selected_board_id': board_id,
        'selected_board_title': board_title,
    }
    return render(request, 'rates.html', context)


class ProjectRateForm(forms.ModelForm):
    rate = forms.IntegerField(required=False, widget=forms.NumberInput(), label="Почасовая ставка")

    class Meta:
        model = ProjectRate
        fields = ('id', 'rate', 'project_id', 'board_id')
        widgets = {
            'id': forms.HiddenInput(),
            'project_id': forms.HiddenInput(),
            'board_id': forms.HiddenInput(),
        }

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.is_default = False
        if self.instance and self.instance.rate is None:
            try:
                dr = DefaultRoleRate.objects.get(role_id=self.instance.role_id)
                if dr.default_rate is not None:
                    self.initial['rate'] = dr.default_rate
                    self.is_default = True
            except DefaultRoleRate.DoesNotExist:
                pass

    def has_changed(self):
        if getattr(self, 'is_default', False):
            return True
        return super().has_changed()

    def clean_rate(self):
        data = self.cleaned_data.get('rate')
        if data == '':
            return None
        return data

def save_rates(request, rates_formset, project_id, project_title, board_id, board_title):
    """
    Валидирует и сохраняет ставки для проекта. Если все поля заполнены, сохраняет formset и перенаправляет с сообщениет об успехе.
    Если какие-либо ставки не заполнены или форма не валидна, возвращает redirect с сообщением об ошибке.
    """
    if rates_formset.is_valid():
        all_filled = all(form.cleaned_data.get('rate') not in (None, '') for form in rates_formset)
        if all_filled:
            rates_formset.save()
            messages.success(request, "Ставки для проекта сохранены.")
        else:
            messages.error(request, "Заполните ставки для всех ролей перед сохранением.")
    else:
        messages.error(request, "Проверьте введённые данные.")
    
    params = {
        'project_id': project_id,
        'project_title': project_title,
        'board_id': board_id,
        'board_title': board_title,
    }
    return redirect(f"/rates/?{urlencode(params)}")


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    # Сбрасывает отступы
    paragraph.paragraph_format.left_indent = 0
    paragraph.paragraph_format.first_line_indent = 0


def replace_placeholder_in_paragraph(paragraph, placeholder, replacement):
    if placeholder in paragraph.text:
        full_text = paragraph.text
        new_text = full_text.replace(placeholder, replacement)
        p = paragraph._element
        for child in list(p):
            p.remove(child)
        paragraph.add_run(new_text)


@login_required
def custom_administration(request):
    if not request.user.is_staff:
        messages.error(request, "У вас нет доступа к странице администрирования, запросите права у администратора.")
        return redirect(request.META.get('HTTP_REFERER', 'rates'))
        
    admin_settings, _ = AdminSettings.objects.get_or_create(pk=1)
    kaiten_roles = fetch_kaiten_roles(
        admin_settings.url_domain_value_id,
        admin_settings.api_auth_key
    )
    kaiten_users = fetch_kaiten_users(
        admin_settings.url_domain_value_id,
        admin_settings.api_auth_key
    )
    overrides = {
        o.kaiten_user_id: o.override_role_id
        for o in KaitenUserRoleOverride.objects.all()
    }
    settings_form = AdminSettingsForm(instance=admin_settings)
    user_form = CustomUserForm()
    
    default_rate_formset = DefaultRoleRateFormSet(queryset=DefaultRoleRate.objects.all())
    
    if request.method == "POST":
        if 'save_user_role' in request.POST:
            uid = request.POST.get('user_id')
            rid = request.POST.get('role_id') or None
            u = User.objects.get(pk=uid)
            u.override_role_id = rid
            u.save()
            messages.success(request, "Роль пользователя сохранена.")
            return redirect('custom_administration')
        if 'save_kaiten_user_role' in request.POST:
            ku_id   = request.POST.get('kaiten_user_id')
            role_id = request.POST.get('role_id') or None
            email = next((u["email"] for u in kaiten_users if str(u["id"]) == ku_id), "")
            KaitenUserRoleOverride.objects.update_or_create(
                kaiten_user_id=ku_id,
                defaults={"email": email, "override_role_id": role_id}
            )
            messages.success(request, "Роль Kaiten-пользователя сохранена.")
            return redirect('custom_administration')
        if 'update_settings' in request.POST:
            settings_form = AdminSettingsForm(request.POST, instance=admin_settings)
            if settings_form.is_valid():
                settings_form.save()
                messages.success(request, "Настройки успешно сохранены.")
                return redirect('custom_administration')
        elif 'create_user' in request.POST:
            user_form = CustomUserForm(request.POST)
            if user_form.is_valid():
                new_user = user_form.save(commit=False)
                password = user_form.cleaned_data.get('password')
                if password:
                    new_user.set_password(password)
                else:
                    new_user.set_password("defaultpassword")
                new_user.save()
                messages.success(request, "Пользователь успешно создан.")
                return redirect('custom_administration')
            else:
                for field, errors in user_form.errors.items():
                    for error in errors:
                        messages.error(request, f"{field.capitalize()}: {error}")
                return redirect('custom_administration')
        elif 'save_default_rates' in request.POST:
            default_rate_formset = DefaultRoleRateFormSet(request.POST, queryset=DefaultRoleRate.objects.all())
            if default_rate_formset.is_valid():
                default_rate_formset.save()
                messages.success(request, "Стандартные ставки сохранены.")
                return redirect('custom_administration')
            else:
                messages.error(request, "Проверьте введённые данные в стандартных ставках.")

    context = {
        'settings_form': settings_form,
        'user_form': user_form,
        'default_rate_formset': default_rate_formset,
        'users': User.objects.all().order_by('id'),
        'kaiten_roles': kaiten_roles,
        'kaiten_users': kaiten_users,
        'overrides':    overrides,
    }
    return render(request, 'custom_administration.html', context)


@login_required
def reports_view(request):
    kaiten_api_down = False
    admin_settings, _ = AdminSettings.objects.get_or_create(pk=1)
    domain = admin_settings.url_domain_value_id
    bearer_key = admin_settings.api_auth_key

    projects = fetch_kaiten_projects(domain, bearer_key)
    if not projects:
        kaiten_api_down = True
    else:
        roles = fetch_kaiten_roles(domain, bearer_key)
        if not roles:
            kaiten_api_down = True
        else:
            for role in roles:
                DefaultRoleRate.objects.get_or_create(
                    role_id=str(role.get('id')),
                    defaults={'role_name': role.get('name')}
                )
            for project in projects:
                pid = str(project.get('id'))
                boards = fetch_kaiten_boards(domain, bearer_key, pid)
                valid = False
                for board in boards:
                    board_roles = fetch_kaiten_board_roles(domain, bearer_key, pid, board['id'])
                    board_valid, _ = check_board_rates(board, pid, board_roles)
                    if board_valid:
                        valid = True
                        break
                project['has_rates'] = valid

    today = datetime.now(pytz.timezone('Europe/Moscow')).strftime('%Y-%m-%d')

    if request.method == "POST":
        project_id = request.POST.get('project')
        board_id = request.POST.get('board')
        template_id = request.POST.get('template')
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')
        selected_project = next((p for p in projects if str(p.get('id')) == project_id), None)

        if selected_project and selected_project.get('has_rates') and board_id and template_id and start_date and end_date:
            try:
                template_instance = TemplateFile.objects.get(id=template_id)
            except TemplateFile.DoesNotExist:
                messages.error(request, "Выбран некорректный шаблон.")
                return redirect('reports')
            return generate_report(request, selected_project, template_instance, start_date, end_date, board_id)

        if not selected_project or not selected_project.get('has_rates'):
            messages.error(request, "Выбран некорректный проект или для проекта не заданы ставки ни для одной доски.")
        elif not board_id:
            messages.error(request, "Выберите доску.")
        elif not template_id:
            messages.error(request, "Выберите шаблон.")
        elif not start_date or not end_date:
            messages.error(request, "Выберите начальную и конечную дату.")

    if kaiten_api_down:
        messages.error(request, "Сервер Kaitен недоступен. Пожалуйста, повторите попытку позже.")

    context = {
        'projects': projects,
        'templates': TemplateFile.objects.all(),
        'today': today,
    }
    return render(request, 'reports.html', context)

@login_required
def generate_report(request, project, template_instance, start_date, end_date, board_id):
    logger.debug("=== START generate_report ===")

    # 1. Настройки и авторизация
    admin_settings, _ = AdminSettings.objects.get_or_create(pk=1)
    domain     = admin_settings.url_domain_value_id
    bearer_key = admin_settings.api_auth_key

    # 2. Определение project_id и project_title
    if isinstance(project, dict):
        project_id    = str(project.get('id', ''))
        project_title = project.get('title', '')
    else:
        # в случае передачи просто ID проекта
        all_projects   = fetch_kaiten_projects(domain, bearer_key)
        proj_dict      = next((p for p in all_projects if str(p.get('id')) == str(project)), {})
        project_id     = str(proj_dict.get('id', ''))
        project_title  = proj_dict.get('title', '')

    # 3. Определение board_id и board_title
    boards     = fetch_kaiten_boards(domain, bearer_key, project_id)
    board_dict = next((b for b in boards if str(b.get('id')) == str(board_id)), {})
    board_id    = str(board_dict.get('id', board_id))
    board_title = board_dict.get('title', '')

    # 4. Загружаем overrides из KaitenUserRoleOverride
    overrides = {
        str(o.kaiten_user_id): o.override_role_id
        for o in KaitenUserRoleOverride.objects.all()
    }
    logger.debug(f"[DEBUG] overrides: {overrides}")

    # 5. Загружаем глобальные роли (для имен)
    roles_list = fetch_kaiten_roles(domain, bearer_key)
    roles_map  = { str(r.get('id')): r.get('name') for r in roles_list }

    # 6. Фетчим карточки и строим таблицу
    # (billing_field_id и billing_field_value должны быть доступны в admin_settings)
    cards = fetch_kaiten_cards(
        domain,
        bearer_key,
        project_id,
        admin_settings.billing_custom_field_id,
        admin_settings.billing_custom_field_value_id
    )

    table_rows   = []
    total_time   = 0.0
    total_amount = 0.0

    for card in cards:
        card_id    = card.get('id')
        card_title = card.get('title')
        logs       = fetch_kaiten_time_logs(domain, bearer_key, card_id)

        for log in logs:
            # Фильтр по дате
            created_iso = log.get('created', '')[:10]
            if not(start_date <= created_iso <= end_date):
                continue

            # Часы
            minutes = float(log.get('time_spent', 0) or 0)
            hours   = minutes / 60.0
            total_time += hours

            # Автор и роль
            author     = log.get('author', {})
            author_id  = str(author.get('id', ''))
            default_role_id = str(log.get('role_id', ''))
            role_id    = overrides.get(author_id, default_role_id)

            # Ставка
            try:
                pr   = ProjectRate.objects.get(
                    project_id=project_id,
                    board_id=board_id,
                    role_id=role_id
                )
                rate = pr.rate or 0
            except ProjectRate.DoesNotExist:
                dr   = DefaultRoleRate.objects.filter(role_id=role_id).first()
                rate = dr.default_rate or 0 if dr else 0
            total_amount += rate * hours

            # Получаем имя роли из Kaiten
            role_name = roles_map.get(role_id, '—')

            # Форматируем строки
            date_str = datetime.strptime(created_iso, '%Y-%m-%d').strftime('%d.%m.%Y')
            table_rows.append({
                'date':       date_str,
                'specialist': author.get('full_name') or author.get('name', '—'),
                'position':   role_name,
                'rate':       f"{rate:.2f}".replace('.', ',') + ' ₽',
                'work':       log.get('comment') or card_title,
                'hours':      f"{hours:.2f}".replace('.', ','),
                'cost':       f"{(rate*hours):.2f}".replace('.', ',') + ' ₽',
            })

    # 7. Проверка наличия данных
    if not table_rows:
        messages.error(request, "Записей не найдено")
        return redirect('reports')

    # 8. Сортировка
    table_rows.sort(key=lambda r: r['date'])

    # 9. Итоги
    total_time_str   = f"{total_time:.2f} ч"
    total_amount_str = (
        f"{total_amount:.2f}".replace('.', ',') +
        ' ₽ (' + convert_number_to_text(total_amount) + ')'
    )

    # 10. Генерация DOCX
    doc = Document(template_instance.file.path)
    # Замена плейсхолдеров
    for p in doc.paragraphs:
        replace_placeholder_in_paragraph(p, '{project_title}', project_title)
        replace_placeholder_in_paragraph(p, '{board_title}', board_title)
        replace_placeholder_in_paragraph(p, '{start_date}', start_date)
        replace_placeholder_in_paragraph(p, '{end_date}', end_date)
        replace_placeholder_in_paragraph(p, '{total_time_spent}',  total_time_str)
        replace_placeholder_in_paragraph(p, '{total_amount_spent}', total_amount_str)

    # Вставка таблицы
    for p in doc.paragraphs:
        if '{table}' in p.text:
            before, after = p.text.split('{table}', 1)
            p.text = before.strip()
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Normal Table'
            set_table_borders(table)
            hdr = table.rows[0].cells
            headers = ['Дата', 'Специалист', 'Позиция', 'Ставка, руб.', 'Содержание работ', 'Часы', 'Стоимость']
            for i, h in enumerate(headers):
                hdr[i].text = h
            insert_table_after(p, table)
            if after.strip():
                insert_paragraph_after_table(table, after.strip())
            for row in table_rows:
                cells = table.add_row().cells
                cells[0].text = row['date'];       cells[0].width = Inches(1.2)
                cells[1].text = row['specialist']; cells[1].width = Inches(2)
                cells[2].text = row['position'];   cells[2].width = Inches(1.5)
                cells[3].text = row['rate'];       cells[3].width = Inches(2.5)
                cells[4].text = row['work'];       cells[4].width = Inches(3.5)
                cells[5].text = row['hours'];      cells[5].width = Inches(1.2)
                cells[6].text = row['cost'];       cells[6].width = Inches(1.5)
            break

    # 11. Стилизация
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing = 1
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(11)
                    if i == 0 or j == 0:
                        run.font.bold = True

    # 12. Отправка документа
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    today = datetime.now(pytz.timezone('Europe/Moscow')).strftime('%Y-%m-%d')
    response['Content-Disposition'] = f'attachment; filename="report_{project_id}_{today}.docx"'
    logger.debug("=== END generate_report ===")
    return response
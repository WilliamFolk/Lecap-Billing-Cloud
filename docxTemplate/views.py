import os
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import TemplateFile
from .forms import TemplateUploadForm, TemplateRenameForm
from docx import Document
from django.core.files import File
from django.core.files.storage import default_storage
from .storage import CustomStorage
custom_storage = CustomStorage()
from num2words import num2words
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

def get_page_count(file_path):
    # Заглушка для проверки количества страниц
    return 1


def insert_table_after(paragraph, table):
    p = paragraph._element
    parent = p.getparent()
    index = list(parent).index(p)
    parent.insert(index + 1, table._element)

def insert_paragraph_after_table(table, text):
    new_p = OxmlElement("w:p")
    table._element.addnext(new_p)
    new_para = Paragraph(new_p, table._parent)
    new_para.text = text
    return new_para

def convert_number_to_text(amount):
    def morph(n, form1, form2, form5):
        n = abs(n) % 100
        if 11 <= n <= 19:
            return form5
        n = n % 10
        if n == 1:
            return form1
        if 2 <= n <= 4:
            return form2
        return form5

    rubles = int(amount)
    kopeks = round((amount - rubles) * 100)
    rubles_text = num2words(rubles, lang='ru')
    kopeks_text = num2words(kopeks, lang='ru', gender='feminine') if kopeks else "ноль"
    
    rubles_declension = morph(rubles, "рубль", "рубля", "рублей")
    kopeks_declension = morph(kopeks, "копейка", "копейки", "копеек")
    
    return f"{rubles_text.capitalize()} {rubles_declension} {kopeks_text} {kopeks_declension}"

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'auto')
        tblBorders.append(element)
    tblPr.append(tblBorders)


@login_required
def upload_template(request):
    templates = TemplateFile.objects.all()
    if request.method == 'POST' and 'file' in request.FILES:
        form = TemplateUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['file']
            if not uploaded_file.name.lower().endswith('.docx'):
                messages.error(request, "Файл должен иметь формат .docx")
                return redirect('docxTemplate:templates')
            if templates.count() >= 100:
                messages.error(request, "Достигнут лимит сохранённых шаблонов (100 файлов)")
                return redirect('docxTemplate:templates')
            temp_path = 'temp_uploaded.docx'
            with open(temp_path, 'wb+') as temp_file:
                for chunk in uploaded_file.chunks():
                    temp_file.write(chunk)
            try:
                doc = Document(temp_path)
            except Exception as e:
                messages.error(request, "Ошибка при открытии файла. Проверьте корректность файла.")
                os.remove(temp_path)
                return redirect('docxTemplate:templates')
            full_text = "\n".join([para.text for para in doc.paragraphs])
            if full_text.count("{table}") != 1:
                messages.error(request, "В файле должен быть ровно один тег {table}")
                os.remove(temp_path)
                return redirect('docxTemplate:templates')
            if get_page_count(temp_path) > 1000:
                messages.error(request, "Файл содержит более 1000 страниц")
                os.remove(temp_path)
                return redirect('docxTemplate:templates')
            with open(temp_path, 'rb') as f:
                django_file = File(f)
                original_name = os.path.basename(uploaded_file.name)
                basename, extension = os.path.splitext(original_name)
                new_name = original_name
                counter = 2
                relative_path = os.path.join('docxTemplate', 'savedFiles', new_name)
                full_path = custom_storage.path(relative_path)
                # Использует os.path.exists для проверки наличия файла по абсолютному пути
                while os.path.exists(full_path):
                    new_name = f"{basename}{counter}{extension}"
                    counter += 1
                    relative_path = os.path.join('docxTemplate', 'savedFiles', new_name)
                    full_path = custom_storage.path(relative_path)
                new_template = TemplateFile()
                new_template.file.save(new_name, django_file)
            os.remove(temp_path)
            messages.success(request, "Шаблон успешно загружен")
            return redirect('docxTemplate:templates')
    else:
        form = TemplateUploadForm()
    return render(request, 'templates.html', {'form': form, 'templates': templates})

def rename_template(request, template_id):
    template_file = get_object_or_404(TemplateFile, id=template_id)
    current_full_name = os.path.basename(template_file.file.name)
    current_base, current_ext = os.path.splitext(current_full_name)

    if request.method == 'POST':
        form = TemplateRenameForm(request.POST)
        if form.is_valid():
            new_input = form.cleaned_data['new_name'].strip()
            new_base = os.path.splitext(new_input)[0]
            if not new_base:
                messages.error(request, "Имя файла не может быть пустым. Имя файла не изменено.")
                return redirect('docxTemplate:templates')

            if new_base == current_base:
                messages.info(request, "Имя файла не изменилось.")
                return redirect('docxTemplate:templates')
            new_full_name = new_base + '.docx'
            # "Django-путь" (с прямыми слешами), который пойдет в БД
            # т. е. docxTemplate/savedFiles/new_file_name.docx
            django_new_path = 'docxTemplate/savedFiles/' + new_full_name
            # Проверка коллизии в хранилище
            counter = 2
            while custom_storage.exists(django_new_path):
                new_full_name = f"{new_base}({counter}).docx"
                django_new_path = 'docxTemplate/savedFiles/' + new_full_name
                counter += 1
            full_old_path = custom_storage.path(template_file.file.name)
            full_new_path = custom_storage.path(django_new_path)
            os.rename(full_old_path, full_new_path)
            template_file.file.name = django_new_path
            template_file.save()
            messages.success(request, "Имя файла изменено")
            return redirect('docxTemplate:templates')
    else:
        form = TemplateRenameForm(initial={'new_name': current_base})
    return render(request, 'rename_template.html', {'form': form, 'template': template_file})


def delete_template(request, template_id):
    template_file = get_object_or_404(TemplateFile, id=template_id)
    template_file.file.delete(save=False)
    template_file.delete()
    messages.success(request, "Шаблон удален")
    return redirect('docxTemplate:templates')

def view_template(request, template_id):
    template_file = get_object_or_404(TemplateFile, id=template_id)
    file_path = default_storage.path(template_file.file.name)
    try:
        doc = Document(file_path)
        content = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        content = "Не удалось прочитать содержимое файла."
    return render(request, 'view_template.html', {'template': template_file, 'content': content})